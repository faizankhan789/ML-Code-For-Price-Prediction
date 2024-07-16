import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, cross_val_score, KFold
from sklearn.ensemble import GradientBoostingRegressor, BaggingRegressor, AdaBoostRegressor
from xgboost import XGBRegressor
from sklearn.svm import SVR
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.pipeline import Pipeline
from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.linear_model import Lasso, Ridge
from skopt import BayesSearchCV
import joblib
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import shap
from docx import Document

# Create a document
doc = Document()
doc.add_heading('Model Evaluation Report', level=1)

# Load and preprocess data
data = pd.read_csv('/content/HyundaiTucson.csv')
data['Mileage(km)'] = data['Mileage(km)'].astype(str).str.replace(',', '').astype(float)
data['Engine Capacity(cc)'] = data['Engine Capacity(cc)'].astype(str).str.replace(',', '').astype(float)
data['Price(PKR lacs)'] = pd.to_numeric(data['Price(PKR lacs)'], errors='coerce')
data.ffill(inplace=True)
data.bfill(inplace=True)

categorical_features = ['Make', 'Model', 'Fuel Type', 'Transmission', 'Registered in', 'Color', 'Assembly', 'Body Type']
numerical_features = ['Model Year', 'Mileage(km)', 'Engine Capacity(cc)']

categorical_transformer = Pipeline(steps=[
    ('imputer', SimpleImputer(strategy='most_frequent')),
    ('onehot', OneHotEncoder(handle_unknown='ignore'))
])

numerical_transformer = Pipeline(steps=[
    ('imputer', SimpleImputer(strategy='median')),
    ('scaler', StandardScaler())
])

preprocessor = ColumnTransformer(
    transformers=[
        ('num', numerical_transformer, numerical_features),
        ('cat', categorical_transformer, categorical_features)
    ])

# Split data into features and target
X = data.drop('Price(PKR lacs)', axis=1)
y = data['Price(PKR lacs)']

# Split the data into training and testing sets and save them to separate files
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
train_data = pd.concat([X_train, y_train], axis=1)
test_data = pd.concat([X_test, y_test], axis=1)

train_data.to_csv('train_data.csv', index=False)
test_data.to_csv('test_data.csv', index=False)

# Define models and their hyperparameter spaces
models = [
    ('GradientBoosting', GradientBoostingRegressor(random_state=42), {
        'regressor__n_estimators': (100, 500),
        'regressor__learning_rate': (0.01, 0.1, 'log-uniform'),
        'regressor__max_depth': (3, 7)
    }),
    ('Bagging', BaggingRegressor(random_state=42), {
        'regressor__n_estimators': (10, 100)
    }),
    ('AdaBoost', AdaBoostRegressor(random_state=42), {
        'regressor__n_estimators': (50, 500),
        'regressor__learning_rate': (0.01, 1.0, 'log-uniform')
    }),
    ('XGBoost', XGBRegressor(random_state=42), {
        'regressor__n_estimators': (100, 500),
        'regressor__learning_rate': (0.01, 0.1, 'log-uniform'),
        'regressor__max_depth': (3, 7)
    }),
    ('SVR', SVR(), {
        'regressor__C': (0.1, 10.0, 'log-uniform'),
        'regressor__epsilon': (0.01, 0.1, 'log-uniform')
    }),
    ('Lasso', Lasso(random_state=42), {
        'regressor__alpha': (0.001, 1.0, 'log-uniform')
    }),
    ('Ridge', Ridge(random_state=42), {
        'regressor__alpha': (0.1, 10.0, 'log-uniform')
    })
]

# Evaluate each model using Bayesian Optimization and K-Fold Cross-Validation
best_models = {}
kf = KFold(n_splits=5, shuffle=True, random_state=42)

for name, model, params in models:
    print(f"Optimizing {name}")
    doc.add_heading(f"Optimizing {name}", level=2)
    pipeline = Pipeline(steps=[('preprocessor', preprocessor), ('regressor', model)])
    opt = BayesSearchCV(estimator=pipeline, search_spaces=params, n_iter=30, cv=kf, n_jobs=-1, scoring='r2', random_state=42)
    opt.fit(X_train, y_train)
    best_models[name] = opt.best_estimator_

    # Display model parameters
    print(f"Best parameters for {name}: {opt.best_params_}")
    doc.add_paragraph(f"Best parameters for {name}: {opt.best_params_}")

    # Evaluate the model on the training set
    y_train_pred = opt.predict(X_train)
    y_test_pred = opt.predict(X_test)

    mse_train = mean_squared_error(y_train, y_train_pred)
    mae_train = mean_absolute_error(y_train, y_train_pred)
    r2_train = r2_score(y_train, y_train_pred)
    mse_test = mean_squared_error(y_test, y_test_pred)
    mae_test = mean_absolute_error(y_test, y_test_pred)
    r2_test = r2_score(y_test, y_test_pred)

    print(f"{name} - Train Mean Squared Error: {mse_train}")
    print(f"{name} - Train Mean Absolute Error: {mae_train}")
    print(f"{name} - Train R^2 Score: {r2_train}")
    print(f"{name} - Test Mean Squared Error: {mse_test}")
    print(f"{name} - Test Mean Absolute Error: {mae_test}")
    print(f"{name} - Test R^2 Score: {r2_test}")

    doc.add_paragraph(f"{name} - Train Mean Squared Error: {mse_train}")
    doc.add_paragraph(f"{name} - Train Mean Absolute Error: {mae_train}")
    doc.add_paragraph(f"{name} - Train R^2 Score: {r2_train}")
    doc.add_paragraph(f"{name} - Test Mean Squared Error: {mse_test}")
    doc.add_paragraph(f"{name} - Test Mean Absolute Error: {mae_test}")
    doc.add_paragraph(f"{name} - Test R^2 Score: {r2_test}")

# Select the best model based on the highest mean cross-validated R^2 score
best_model_name = max(best_models, key=lambda name: cross_val_score(best_models[name], X_train, y_train, cv=kf, scoring='r2').mean())
best_model = best_models[best_model_name]

# Evaluate the best model on the test set
y_pred = best_model.predict(X_test)
mse_best = mean_squared_error(y_test, y_pred)
mae_best = mean_absolute_error(y_test, y_pred)
r2_best = r2_score(y_test, y_pred)

print(f"Best model: {best_model_name}")
print(f"Mean Squared Error: {mse_best}")
print(f"Mean Absolute Error: {mae_best}")
print(f"R^2 Score: {r2_best}")

doc.add_heading(f"Best model: {best_model_name}", level=2)
doc.add_paragraph(f"Mean Squared Error: {mse_best}")
doc.add_paragraph(f"Mean Absolute Error: {mae_best}")
doc.add_paragraph(f"R^2 Score: {r2_best}")

# Save the best model
joblib.dump(best_model, f'best_car_price_predictor_{best_model_name}.pkl')

# Get feature names after one-hot encoding
ohe = best_model['preprocessor'].named_transformers_['cat']['onehot']
ohe_feature_names = ohe.get_feature_names_out(categorical_features)

# Combine numerical and encoded categorical feature names
feature_names = np.concatenate([numerical_features, ohe_feature_names])

# Model interpretation with SHAP
# Convert the transformed data to a dense format
X_train_transformed = best_model['preprocessor'].transform(X_train).toarray()
X_test_transformed = best_model['preprocessor'].transform(X_test).toarray()

# Initialize the SHAP explainer with the regressor part of the best model
explainer = shap.Explainer(best_model['regressor'], X_train_transformed)


# Calculate SHAP values
shap_values = explainer(X_test_transformed)

# Generate SHAP summary plot
shap.summary_plot(shap_values, X_test_transformed, feature_names=feature_names)

# Function to get the formula and explanation for each model
def get_model_formula_and_explanation(model):
    if isinstance(model['regressor'], GradientBoostingRegressor):
        formula = "F(x) = ∑[m=1 to M] γ_m h_m(x)"
        explanation = """
        Gradient Boosting Regressor:
        - Formula: F(x) = ∑[m=1 to M] γ_m h_m(x)
          where F(x) is the final prediction,
          M is the number of boosting stages (trees),
          γ_m is the learning rate applied to the m-th tree,
          and h_m(x) is the m-th tree's prediction.
        - Parameters:
          - n_estimators: Number of boosting stages (trees) to be run.
          - learning_rate: Reduces the contribution of each tree to prevent overfitting.
          - max_depth: Limits the maximum depth of each tree, controlling model complexity.
        """
    elif isinstance(model['regressor'], BaggingRegressor):
        formula = "F(x) = (1/n) ∑[i=1 to n] h_i(x)"
        explanation = """
        Bagging Regressor:
        - Formula: F(x) = (1/n) ∑[i=1 to n] h_i(x)
          where F(x) is the final prediction,
          n is the number of base estimators (trees),
          and h_i(x) is the prediction of the i-th tree.
        - Parameters:
          - n_estimators: Number of base estimators (trees) to be averaged.
        """
    elif isinstance(model['regressor'], AdaBoostRegressor):
        formula = "F(x) = ∑[m=1 to M] α_m h_m(x)"
        explanation = """
        AdaBoost Regressor:
        - Formula: F(x) = ∑[m=1 to M] α_m h_m(x)
          where F(x) is the final prediction,
          M is the number of boosting stages (weak learners),
          α_m is the weight of the m-th learner,
          and h_m(x) is the m-th learner's prediction.
        - Parameters:
          - n_estimators: Number of boosting stages.
          - learning_rate: Shrinks the contribution of each weak learner.
        """
    elif isinstance(model['regressor'], XGBRegressor):
        formula = "F(x) = ∑[k=1 to K] γ_k h_k(x)"
        explanation = """
        XGBoost Regressor:
        - Formula: F(x) = ∑[k=1 to K] γ_k h_k(x)
          where F(x) is the final prediction,
          K is the number of boosting rounds (trees),
          γ_k is the learning rate applied to the k-th tree,
          and h_k(x) is the k-th tree's prediction.
        - Parameters:
          - n_estimators: Number of boosting rounds.
          - learning_rate: Shrinks the contribution of each tree to prevent overfitting.
          - max_depth: Limits the maximum depth of each tree.
        """
    elif isinstance(model['regressor'], SVR):
        formula = "f(x) = ⟨w,x⟩ + b"
        explanation = """
        SVR (Support Vector Regressor):
        - Formula: f(x) = ⟨w,x⟩ + b
          where f(x) is the predicted value,
          ⟨w,x⟩ is the dot product of the weight vector w and input x,
          and b is the bias term.
        - Parameters:
          - C: Regularization parameter that balances margin maximization and error minimization.
          - epsilon: Specifies the epsilon-tube within which no penalty is associated in the training loss function.
        """
    elif isinstance(model['regressor'], Lasso):
        formula = "min_w { (1/2n) ∑[i=1 to n] (y_i - ⟨w, x_i⟩)^2 + α‖w‖_1 }"
        explanation = """
        Lasso Regressor:
        - Formula: min_w { (1/2n) ∑[i=1 to n] (y_i - ⟨w, x_i⟩)^2 + α‖w‖_1 }
          where (1/2n) ∑[i=1 to n] (y_i - ⟨w, x_i⟩)^2 is the least squares loss,
          and α‖w‖_1 is the L1 regularization term, with α as the regularization strength.
        - Parameters:
          - alpha: Regularization parameter that controls the strength of the L1 penalty.
        """
    elif isinstance(model['regressor'], Ridge):
        formula = "min_w { (1/2n) ∑[i=1 to n] (y_i - ⟨w, x_i⟩)^2 + α‖w‖_2^2 }"
        explanation = """
        Ridge Regressor:
        - Formula: min_w { (1/2n) ∑[i=1 to n] (y_i - ⟨w, x_i⟩)^2 + α‖w‖_2^2 }
          where (1/2n) ∑[i=1 to n] (y_i - ⟨w, x_i⟩)^2 is the least squares loss,
          and α‖w‖_2^2 is the L2 regularization term, with α as the regularization strength.
        - Parameters:
          - alpha: Regularization parameter that controls the strength of the L2 penalty.
        """
    else:
        formula = "Unknown"
        explanation = "Unknown model"

    return formula, explanation

# Display the formula and explanation for the best model
formula, explanation = get_model_formula_and_explanation(best_model)
doc.add_heading(f"Formula and Explanation for {best_model_name}", level=2)
doc.add_paragraph(f"Formula: {formula}")
doc.add_paragraph(f"Explanation: {explanation}")

print(f"Formula and Explanation for {best_model_name}")
print(f"Formula: {formula}")
print(f"Explanation: {explanation}")

# Save the document
doc.save('model_evaluation_report.docx')

print("Model evaluation report saved as 'model_evaluation_report.docx'")
